# -*- coding: utf-8 -*-
from tkinter import messagebox
from tkinter.filedialog import asksaveasfilename
from anytree import Node
from anytree.exporter import DotExporter
from docx import Document
from docx.shared import Cm
from openpyxl import load_workbook
from pulp import *
from math import ceil, floor
import string
from sympy import Symbol, integrate
import platform
import os
import matplotlib.pyplot as plt
import numpy as np

'''глобальные переменные'''
# чтобы прописать graphviz в PATH
os.environ["PATH"] += os.pathsep + r'C:\Program Files (x86)\Graphviz2.38\bin'
# допустимые символы в массивах excel
expr_allowed_symbols = set(string.digits + '.' + '*' + '/' + '(' + ')' + 'x' + 'X' + '+' + '-')
symbols = set('.' + '*' + '/' + '(' + ')' + '+' + '-')
results_dir = 'results'
if not os.path.exists(results_dir):
    os.mkdir(results_dir)


class Solved(object):

    def __init__(self, problem, number, func_value, vars_value, status,
                 cont_var=None, cont_var_value=None, parent_number=None):
        self.problem = problem
        self.status = status
        self.number = number
        self.func_value = func_value
        self.vars_value = vars_value
        self.cont_var = cont_var
        self.cont_var_value = cont_var_value
        self.parent_number = parent_number

    def __repr__(self):
        return str(self.number)

    # чтобы решить проблему с pyinstaller
    if platform.system() == 'Windows':
        solver = COIN_CMD(path=os.path.join(os.getcwd(), 'solver\\win\\cbc.exe'))
    else:
        solver = COIN_CMD(path=os.path.join(os.getcwd(), 'solver/linux/cbc'))
    tree = []
    statuses = ['Не решено', 'Целое', 'Неопределенно', 'Не ограниченно', 'Нецелое']

    def nodenamefunc(node):
        node_params = [str(node.name), node.status, str(node.func_value), *node.xs]
        return "\n".join(node_params)

    def make_node(self):
        status = Solved.statuses[self.status]
        if self.status == 1 and self.cont_var is not None:
            status = 'Нецелочисленное'
        xs = [str(x[0]) + ' = ' + str(x[1]) for x in zip(self.problem.variables(), self.vars_value)]
        new_node = Node(name=self.number, status=status, xs=xs,
                        func_value=self.func_value, parent_name=self.parent_number)
        for node in self.tree:
            if node.name == new_node.parent_name:
                new_node.parent = node
        self.tree.append(new_node)


class Solution(object):
    def __init__(self, acc, solution=None, optimal_problems=[], auto_coeff_D=False):
        self.acc = acc
        self.coeff_D = auto_coeff_D
        self.has_sol = False
        if solution is not None:
            for x in solution.vars_value:
                if x != 0:
                    self.has_sol = True
                    self.optimal_problems = optimal_problems
                    self.variables = solution.problem.variables()
                    self.func_value = solution.func_value
                    self.vars_value = solution.vars_value
                    self.number = solution.number
                    break

    def __repr__(self):
        return str(self.number)


def create_Solved(problem, acc, parent_number=None):
    problem_copy = problem.deepcopy()
    problem_copy.solve(Solved.solver)
    acc += 1
    # создаем объект решенной задачи
    solved = Solved(problem=problem,
                    status=problem_copy.status,
                    number=acc,
                    func_value=value(problem_copy.objective),
                    vars_value=[var.varValue for var in problem_copy.variables()],
                    parent_number=parent_number)
    for v in problem_copy.variables():
        if v.varValue != int(v.varValue):
            solved.cont_var, solved.cont_var_value = v, v.varValue
            break
    solved.make_node()
    return solved, acc


def get_values(worksheet, row, column, column_len=-1, is_expression=False, is_name=False):
    values = []
    allowed = expr_allowed_symbols
    while worksheet.cell(row, column).value is not None:
        #  если это выражение, все его символы разрешены и оно состоит не из одних лишь операторов
        cell_value = worksheet.cell(row, column).value
        if is_name:
            values.append(str(cell_value))
            row += 1
        elif is_expression and set(str(cell_value)) <= allowed and not set(str(cell_value)) <= symbols:
            values.append(str(cell_value))
            row += 1
        elif (type(cell_value) is float) or (type(cell_value) is int):
            values.append(cell_value)
            row += 1
        else:
            raise Exception(
                messagebox.showinfo('Ошибка',
                                    'В листе Excel в ячейке (' + str(column) + ', ' + str(row) + ') '
                                    'значение написано неверно или использованы недопустимые символы. '
                                    '\nДопустимые: ' + ' '.join(sorted(allowed)))
            )
    if column_len != -1 and len(values) != column_len:
        raise Exception(
            messagebox.showinfo('Ошибка', 'В листе Excel кол-во элементов в колонке '
                                + str(column) + ' не равно кол-ву элементов в предшествующей'))
    return values  # не прерывается


def get_data_excel(filepath, T):
    #  открываем нужный лист в выбранной в форме книге excel
    workbook = load_workbook(filename=str(filepath))
    worksheet = workbook.worksheets[0]
    #  получаем массивы данных с листа по столбцам
    name = get_values(worksheet, 2, 1, is_name=True)
    alpha = get_values(worksheet, 2, 2, len(name))
    beta = get_values(worksheet, 2, 3, len(name))
    v = get_values(worksheet, 2, 4, len(name))
    V = get_values(worksheet, 2, 5, len(name))
    teta = get_values(worksheet, 2, 6, len(name), is_expression=True)
    # добавляем первоначальный порядок
    order = [i for i in range(0, len(name))]
    # интегрируем тета
    x = Symbol('x')
    try:
        teta_integrated = [float(integrate(eval(teta[i]), (x, 0, T))) for i in range(0, len(teta))]
    except Exception:
        messagebox.showinfo('Ошибка', 'Интенсивности реализации товара заданы некорректно')
        raise
    data = {'name': name,
            'alpha': alpha,
            'beta': beta,
            'v': v,
            'unsorted_v': v,
            'teta_integrated': teta_integrated,
            'V': V,
            'order': order}
    return data, workbook, worksheet


def sort_data(sort_type, data):
    if sort_type == 'b/a':
        ba = [b / a for a, b in zip(data['alpha'], data['beta'])]
        zipped = list(zip(ba, data['teta_integrated'], data['alpha'],
                          data['beta'], data['V'], data['v'], data['order']))
        zipped.sort(reverse=True)
        ba, data['teta_integrated'], data['alpha'], data['beta'], \
            data['V'], data['v'], data['order'] = zip(*zipped)
    elif sort_type == 'teta':
        zipped = list(zip(data['teta_integrated'], data['alpha'],
                          data['beta'], data['V'], data['v'], data['order']))
        zipped.sort(reverse=True)
        data['teta_integrated'], data['alpha'], data['beta'], \
            data['V'], data['v'], data['order'] = zip(*zipped)
    return data


def form_problem(data, **coeffs):

    def form_problem_1(alpha, beta, v, teta_integrated, V, F):
        k = [a / b for a, b in zip(V, v)]
        n = len(alpha)
        problem = LpProblem('Zadachka', LpMaximize)
        x = LpVariable.dicts('x', range(n), lowBound=0, cat=LpContinuous)
        sum_var1 = lpSum([x[i] * v[i] * beta[i] for i in range(0, n)])
        sum_var2 = lpSum([x[i] * v[i] * alpha[i] for i in range(0, n)])
        problem += sum_var1 + (F - sum_var2)  # 'Функция цели "11.1"'
        problem += sum_var2 <= F  # "11.2"
        constraint1 = [x[i] <= k[i] for i in range(0, n)]
        for cnstr in constraint1:
            problem += cnstr
        constraint2 = [x[i] * v[i] <= min(teta_integrated[i], V[i]) for i in range(0, n)]
        for cnstr in constraint2:
            problem += cnstr
        constraint3 = [min(teta_integrated[i], V[i]) <= v[i] * (x[i] + 1) for i in range(0, n)]
        for cnstr in constraint3:
            problem += cnstr
        return problem

    def form_problem_2(alpha, beta, v, teta_integrated, V, F, D, y):
        k = [a / b for a, b in zip(V, v)]
        n = len(alpha)
        V = [v[i] * k[i] for i in range(0, n)]
        problem = LpProblem('Zadachka', LpMaximize)
        x = LpVariable.dicts('x', range(n), lowBound=0, cat=LpContinuous)
        sum_xvb = lpSum([x[i] * v[i] * beta[i] for i in range(0, n)])
        sum_xva = lpSum([x[i] * v[i] * alpha[i] for i in range(0, n)])
        sum_1yax = lpSum([(1 + y) * alpha[i] * x[i] for i in range(0, n)])
        sum_bx = lpSum([beta[i] * x[i] for i in range(0, n)])
        problem += sum_xvb + ((1 + y) * (F - sum_xva))  # цель
        problem += sum_xva <= F + D
        constraint1 = [x[i] <= k[i] for i in range(0, n)]
        constraint2 = [x[i] * v[i] <= min(teta_integrated[i], V[i]) for i in range(0, n)]
        constraint3 = [min(teta_integrated[i], V[i]) <= v[i] * (x[i] + 1) for i in range(0, n)]
        for cnstr in [*constraint1, *constraint2, *constraint3]:
            problem += cnstr
        problem += sum_1yax <= sum_bx
        return problem

    alpha = data['alpha']
    beta = data['beta']
    v = data['v']
    V = data['V']
    teta_integrated = data['teta_integrated']
    if coeffs['zadacha'] == 1:
        return form_problem_1(alpha, beta, v, teta_integrated, V, coeffs['F'])
    elif coeffs['zadacha'] == 2:
        return form_problem_2(alpha, beta, v, teta_integrated, V, coeffs['F'], coeffs['D'], coeffs['y'])


def solve_problem(problem):
    # обнуляем дерево, потому что в одном прогоне программы может решаться несколько задач
    Solved.tree = []
    optimal = []  # лист оптимальных целочисленных решений
    queue = []  # очередь на ветвление
    acc = 0  # счетчик задач
    max_z = 0  # максимальное значение целевой функции

    first, acc = create_Solved(problem, acc, parent_number=None)
    # возвращаем задачу, если она не оптимальна
    if first.status != 1:
        return Solution(acc=acc, solution=None)
    else:
        # проверяем проблему на целочисленность
        if first.cont_var is not None:
            # добавляем первую проблему в очередь на ветвление
            queue.append(first)
            return branch_and_bound(queue, max_z, acc, optimal)
        # если проблема целочислена на первом шаге
        else:
            return Solution(acc=acc, solution=first, optimal_problems=[first])


def branch_and_bound(queue, max_z, acc, optimal):  # передаем сюда задачу на ветвление, в том числе нецелую переменную
    if queue:
        max_prob = queue[0]
        for prob in queue[1:]:
            if prob.func_value > max_prob.func_value:
                max_prob = prob
        queue.remove(max_prob)
        i = 1
        queue, max_z, acc, optimal = make_branch(max_prob, acc, queue, max_z, optimal, i)
        return branch_and_bound(queue, max_z, acc, optimal)
    else:
        if optimal:
            solution = optimal[0]
            for prob in optimal[1:]:
                if prob.func_value > solution.func_value:
                    solution = prob
            return Solution(acc=acc, solution=solution, optimal_problems=optimal)
        else:
            return Solution(acc=acc, solution=None)


def make_branch(parent_problem, acc, queue, max_z, optimal, i):
    child_problem = parent_problem.problem.deepcopy()
    if i == 1:
        child_problem += parent_problem.cont_var <= floor(parent_problem.cont_var_value)  # левая ветвь
    elif i == 2:
        child_problem += parent_problem.cont_var >= ceil(parent_problem.cont_var_value)  # правая ветвь
    child_solved, acc = create_Solved(child_problem, acc, parent_number=parent_problem.number)
    if child_solved.status == 1:  # формируем подпроблему в очередь
        if child_solved.cont_var is None and max_z == 0:  # шаг 5
            max_z = child_solved.func_value
            optimal.append(child_solved)
            for prob in queue:  # проверяем, будем ли ветвить нецелые задачи, оставшиеся в queue
                if prob.func_value < max_z:
                    queue.remove(prob)

        elif child_solved.cont_var is None and child_solved.func_value >= max_z:
            optimal.append(child_solved)

        elif child_solved.cont_var is not None and max_z == 0:
            queue.append(child_solved)

        elif child_solved.cont_var is not None and child_solved.func_value > max_z:
            queue.append(child_solved)

    if i == 1:
        return make_branch(parent_problem, acc, queue, max_z, optimal, i + 1)  # ветвим второй раз
    elif i == 2:
        return queue, max_z, acc, optimal


def solution_stability(solution, data, coeffs):

    def minimum_e(optimal_list, e_min_problem, e_min_prob_index, es):
        """
        Эта функция находит интервалы значений инфляции (е), в рамках которых
        оптимальный портфель закупок сохраняется. Первый шаг расчета начинается с
        оптимальной проблемы, у которой е = 0.
        """
        e_min = float('inf')
        x_l = e_min_problem.vars_value
        l = e_min_prob_index
        n = len(optimal_list)
        for i in range(l + 1, n):
            x = optimal_list[i].vars_value
            e = (
                    (
                            sum(x_l[i] * (data['beta'][i] - data['alpha'][i]) for i in range(0, len(x))) -
                            sum(x[i] * (data['beta'][i] - data['alpha'][i]) for i in range(0, len(x)))
                    ) /
                    (
                            sum(x[i] * data['beta'][i] for i in range(0, len(x))) -
                            sum(x_l[i] * data['beta'][i] for i in range(0, len(x)))
                    )
            )
            if 0 < e < e_min:
                e_min = e
                e_min_problem = optimal_list[i]
                e_min_prob_index = i
        if l != n - 1 and e_min != float('inf'):
            es[e_min] = e_min_problem
            return minimum_e(optimal_list, e_min_problem, e_min_prob_index, es)
        else:
            return es

    def make_stability_plot(es, data, coeffs):
        # нужно построить функцию от неизвестного e, имея две точки: значение функции при e = 0
        # и при e, пересекающем оптимальную функцию. По сути, нужно только посчитать
        # значение целевой функции в этой точке и построить графики. Нужно использовать
        # проблемы l+1
        alpha = data['alpha']
        v = data['v']
        F = coeffs['F']
        n = len(alpha)

        x_opt = es[0].vars_value
        opt_number = es[0].number
        beta = [b * 2 for b in data['beta']]
        sum_var1 = sum([x_opt[i] * v[i] * beta[i] for i in range(0, n)])
        sum_var2 = sum([x_opt[i] * v[i] * alpha[i] for i in range(0, n)])
        opt_zero_f_val = es[0].func_value
        opt_one_f_val = sum_var1 + F - sum_var2
        fig, ax = plt.subplots(nrows=1, ncols=1)
        ax.plot([0, 1], [opt_zero_f_val, opt_one_f_val], label='Оптимальная задача: ' + str(opt_number))
        del es[0]
        if es:
            intersection_f_list = []
            zero_f_list = [problem.func_value for problem in es.values()]
            for e, problem in es.items():
                x = problem.vars_value
                beta = [b * (1 + e) for b in data['beta']]
                sum_var1 = sum([x[i] * v[i] * beta[i] for i in range(0, n)])
                sum_var2 = sum([x[i] * v[i] * alpha[i] for i in range(0, n)])
                intersection_f_list.append(sum_var1 + F - sum_var2)
            for zero_f, cross_f, e, problem in zip(zero_f_list, intersection_f_list, es, es.values()):
                ax.plot([0, e, 2*e], [zero_f, cross_f, cross_f + zero_f],
                        label='Целочисленная задача: ' + str(problem.number))
                ax.plot([e, e], [0, cross_f], ':')

        ax.set_ylabel('Значение целевой функции от инфляции, F(e)')
        ax.set_xlabel('Инфляция в долях, e')
        ax.set_title('Устойчивость решения')
        ax.legend()
        ax.grid()
        plot_path = results_dir + '/temp_plot.png'
        fig.savefig(plot_path, bbox_inches='tight')
        return plot_path, es

    optimal_list = solution.optimal_problems
    if 'dummy' in coeffs and coeffs['dummy']:
        vars_value = solution.vars_value
        func_value = solution.func_value
        optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                                   number='a', func_value=func_value - 0.2 * func_value,
                                   vars_value=[i - 1 if i % 2 == 0 else i + 1 for i in vars_value],
                                   status=1))
        optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                                   number='b', func_value=func_value - 0.4 * func_value,
                                   vars_value=[i - 0.5 if i % 2 == 0 else i + 0.5 for i in vars_value],
                                   status=1))
        optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                                   number='c', func_value=func_value - 0.4 * func_value,
                                   vars_value=[i - 0.8 if i % 2 == 0 else i + 0.8 for i in vars_value],
                                   status=1))
        optimal_list.append(Solved(problem=solution.optimal_problems[0].problem,
                                   number='d', func_value=func_value - 0.4 * func_value,
                                   vars_value=[i - 1.2 if i % 2 == 0 else i + 1.2 for i in vars_value],
                                   status=1))
    k_list = {problem.number: sum(problem.vars_value[i] * data['beta'][i]
            for i in range(0, len(data['beta']))) for problem in optimal_list}
    if coeffs['zadacha'] == 1:
        optimal_list.sort(key=lambda problem: sum(problem.vars_value[i] * data['beta'][i]
                                                  for i in range(0, len(data['beta']))))
    elif coeffs['zadacha'] == 2:
        pass
    sol_num = solution.number
    index_list = [x.number for x in optimal_list]
    sol_index = index_list.index(sol_num) if sol_num in index_list else None
    es = {0: solution}
    if len(optimal_list) > 1:
        es = minimum_e(optimal_list, solution, sol_index, es)
    return make_stability_plot(es, data, coeffs)


def show_results(sort_type, solved, sorted_data, coeffs):
    unsorted_xs = None
    if not solved.has_sol:
        status = 'Статус: Нерешаемо'
        acc = 'Кол-во решенных ЗЛП: ' + str(solved.acc)
        results = [status, acc]
        if solved.coeff_D:
            results.append('Подобранный D: ' + str(solved.coeff_D))
        elif 'D' in coeffs:
            results.append('D: ' + str(coeffs['D']))

    else:
        xs_order = sorted_data['order']
        status = 'Статус: Оптимально'
        unsorted_xs = [0 for i in range(0, len(xs_order))]
        for i in range(0, len(xs_order)):
            unsorted_xs[xs_order[i]] = int(solved.vars_value[i])
        partys = 'Количество партий товаров:'
        xs = [x[0] + ' = ' + str(x[1]) + ' по ' + str(x[2]) + ' штук'
              for x in zip(sorted_data['name'], unsorted_xs, sorted_data['unsorted_v'])]
        number_of_optimal = 'Номер оптимальной задачи: ' + str(solved.number)
        func_value = 'Чистая прибыль: ' + str(solved.func_value - coeffs['F'])
        sort_type = 'Сортировка: ' + sort_type
        acc = 'Кол-во решенных ЗЛП: ' + str(solved.acc)
        results = [status, func_value, partys, *xs, sort_type, number_of_optimal, acc]
        if solved.coeff_D:
            results.append('Подобранный D: ' + str(solved.coeff_D))
        elif 'D' in coeffs:
            results.append('D: ' + str(coeffs['D']))
    tree_img_path = results_dir + '/temp_tree.png'

    DotExporter(Solved.tree[0], nodenamefunc=Solved.nodenamefunc).to_picture(tree_img_path)
    return results, tree_img_path, unsorted_xs


def write_to_excel(workbook, worksheet, filepath, results):
    for row in worksheet['g2:g50']:
        for cell in row:
            cell.value = None
    if results is not None:
        for i in range(0, len(results)):
            worksheet.cell(i + 2, 7).value = results[i]
    while True:
        try:
            workbook.save(filepath)
            break
        except PermissionError:
            if not messagebox.askokcancel('Ошибка', 'Нет доступа к указанному файлу Excel.\n'
                                          'Возможно, Вы не закрыли этот файл.\n'
                                          'Закройте его и нажмите ОК.'):
                break


def write_to_docx(problem, results, plot_img_path, tree_img_path, is_stable, es):
    file_name = asksaveasfilename(defaultextension='.docx',
                                  initialdir='results',
                                  filetypes=[('Word Document (.docx)', '.docx')])
    if file_name:
        document = Document()
        while True:
            try:
                document.save(file_name)
                # создаем кортеж values of OrderedDict, преобразуем в лист и индексируем
                constrs = map(str, list(problem.constraints.values()))
                obj = problem.objective
                document.add_heading('Максимизируем целевую функцию:', 1)
                document.add_paragraph(str(obj) + '\n\nС ограничениями:\n' +
                                       '\n'.join(constrs))
                document.add_paragraph("\n".join(results))
                document.add_heading('Дерево решений задачи:', 1)
                document.add_picture(tree_img_path, width=Cm(12))
                if is_stable:
                    document.add_heading('Устойчивость решения:', 1)
                    if plot_img_path is not None:
                        document.add_picture(plot_img_path, width=Cm(14))
                        document.add_paragraph('Интервалы сохранения значения оптимального портфеля закупок:\n' +
                                               '[0, ' + ''.join(str(e) + ', ' for e in es) + '∞)')
                    else:
                        document.add_paragraph('Задача не решаема')
                document.save(file_name)
                messagebox.showinfo('Файл создан', 'Файл с результатом был успешно сохранен!')
                os.startfile(file_name)
                break
            except PermissionError:
                if not messagebox.askokcancel('Ошибка', 'Доступ к выбранному файлу невозможен. Может быть, '
                                              'Вы пытаетесь переписать открытый файл.\n'
                                              'Попробуйте закрыть его и нажмите ОК'):
                    break


def integer_lp(filepath, **coeffs):
    plot_img_path = None
    is_stable = False
    data, workbook, worksheet = get_data_excel(filepath, coeffs['T'])
    sorted_data = sort_data(coeffs['sort'], data)
    es = None
    # если автоподбор параметра D
    if 'auto_D' in coeffs:
        # решаем проблему с первым значением D и убираем его из списка
        coeffs['D'] = coeffs['auto_D'].pop(0)
        problem = form_problem(sorted_data, **coeffs)
        solution_problem = solve_problem(problem)
        solution_problem.coeff_D = coeffs['D']
        if not solution_problem.has_sol:
            solution_problem.func_value = 0
        for d in coeffs['auto_D']:
            coeffs['D'] = d
            problem = form_problem(sorted_data, **coeffs)
            solved_problem = solve_problem(problem)
            if solved_problem.has_sol and solved_problem.func_value > solution_problem.func_value:
                solution_problem = solved_problem
                solution_problem.coeff_D = d
    else:
        problem = form_problem(sorted_data, **coeffs)
        solution_problem = solve_problem(problem)
    # если стабильность решения и True
    if 'stable' in coeffs and coeffs['stable']:
        is_stable = True
        if solution_problem.has_sol:
            plot_img_path, es = solution_stability(solution_problem, sorted_data, coeffs)
    results, tree_img_path, unsorted_xs = \
        show_results(coeffs['sort'], solution_problem, sorted_data, coeffs)
    write_to_excel(workbook, worksheet, filepath, unsorted_xs)
    answer = messagebox.askyesno('Решение', "\n".join(results) +
                                 '\n\nХотите сохранить подробный результат\nв DOCX файл?')
    if answer:
        write_to_docx(problem, results, plot_img_path, tree_img_path, is_stable, es)
    # удаляю временные файлы, которые закидываются в doc с результатами, графика может не быть
    os.remove(tree_img_path)
    try:
        os.remove(plot_img_path)
    except TypeError:
        pass


def main():
    print(integer_lp('excel/Zadachka.xlsx', T=1, F=30000, zadacha=1, sort='b/a', stable=1))


if __name__ == '__main__':
    main()
